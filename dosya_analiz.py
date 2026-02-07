#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Proje Klasör Analiz Scripti
============================
Yazılımcılara verilen proje klasörlerini analiz eder.
- Klasör yapısını çıkarır
- Excel dosyalarını analiz eder (sayfalar, tablolar, hücre bağlantıları, formüller)
- Word dosyalarını okur ve özetler
- PDF dosylarını okur ve özetler
- Her klasöre MD rapor dosyası oluşturur
- Ana rapor klasörü oluşturur ve tüm raporları toplar

Kullanım:
    python dosya_analiz.py <hedef_klasör_yolu>
    python dosya_analiz.py .   (mevcut klasörü analiz eder)
"""

import os
import sys
import re
import datetime
import shutil
import argparse
from pathlib import Path
from collections import defaultdict

# ─── Opsiyonel kütüphaneler ───────────────────────────────────────────────────
try:
    import openpyxl
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False


# ─── Yardımcı sabitler ────────────────────────────────────────────────────────
EXCEL_EXTENSIONS = {'.xlsx', '.xlsm', '.xltx', '.xltm'}
EXCEL_OLD_EXTENSIONS = {'.xls'}
WORD_EXTENSIONS = {'.docx'}
PDF_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.ico', '.webp', '.tiff'}
CODE_EXTENSIONS = {
    '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.cs', '.cpp', '.c', '.h',
    '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.r', '.m',
    '.html', '.css', '.scss', '.less', '.vue', '.svelte',
    '.sql', '.sh', '.bash', '.zsh', '.ps1', '.bat', '.cmd',
    '.json', '.yaml', '.yml', '.toml', '.xml', '.ini', '.cfg', '.env',
    '.md', '.markdown', '.txt', '.rst', '.tex',
    '.dockerfile', '.dockerignore', '.gitignore', '.editorconfig',
}
ARCHIVE_EXTENSIONS = {'.zip', '.tar', '.gz', '.rar', '.7z', '.bz2', '.xz'}

IGNORED_DIRS = {
    '__pycache__', 'node_modules', '.git', '.svn', '.hg', '.DS_Store',
    '.idea', '.vscode', 'venv', 'env', '.env', '.venv', 'dist', 'build',
    '__MACOSX', '.tox', '.mypy_cache', '.pytest_cache',
}

RAPOR_KLASOR_ADI = "_ANALIZ_RAPORLARI"
RAPOR_DOSYA_ADI = "_KLASOR_RAPORU.md"

# ─── Boyut formatlama ─────────────────────────────────────────────────────────
def format_size(size_bytes):
    """Byte cinsinden boyutu okunabilir formata çevirir."""
    if size_bytes == 0:
        return "0 B"
    units = ['B', 'KB', 'MB', 'GB', 'TB']
    i = 0
    size = float(size_bytes)
    while size >= 1024 and i < len(units) - 1:
        size /= 1024
        i += 1
    return f"{size:.1f} {units[i]}"


# ═══════════════════════════════════════════════════════════════════════════════
# EXCEL ANALİZİ
# ═══════════════════════════════════════════════════════════════════════════════
def analyze_excel(filepath):
    """Excel dosyasını detaylı analiz eder."""
    report = []
    ext = Path(filepath).suffix.lower()

    if ext in EXCEL_OLD_EXTENSIONS:
        return _analyze_excel_xls(filepath)

    if not HAS_OPENPYXL:
        report.append("  > ⚠️ `openpyxl` kütüphanesi yüklü değil. Excel analizi yapılamadı.")
        report.append("  > Yüklemek için: `pip install openpyxl`")
        return "\n".join(report)

    try:
        wb = openpyxl.load_workbook(filepath, data_only=False, read_only=False)
    except Exception as e:
        report.append(f"  > ❌ Dosya açılamadı: {e}")
        return "\n".join(report)

    report.append(f"  - **Sayfa Sayısı:** {len(wb.sheetnames)}")
    report.append(f"  - **Sayfalar:** {', '.join(wb.sheetnames)}")
    report.append("")

    # Tanımlı isimler (Named Ranges)
    try:
        defined_names = list(wb.defined_names.values()) if hasattr(wb.defined_names, 'values') else []
        if defined_names:
            report.append("  #### 📌 Tanımlı İsimler (Named Ranges)")
            report.append("  | İsim | Referans |")
            report.append("  |------|----------|")
            for name in defined_names:
                try:
                    name_str = name.name if hasattr(name, 'name') else str(name)
                    ref_str = name.attr_text if hasattr(name, 'attr_text') else str(name.value) if hasattr(name, 'value') else ''
                    report.append(f"  | `{name_str}` | `{ref_str}` |")
                except:
                    pass
            report.append("")
    except Exception:
        pass

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        report.append(f"  #### 📄 Sayfa: `{sheet_name}`")

        # Boyut bilgisi
        if ws.dimensions and ws.dimensions != 'A1:A1':
            report.append(f"  - **Veri Aralığı:** `{ws.dimensions}`")
        
        min_row = ws.min_row or 1
        max_row = ws.max_row or 1
        min_col = ws.min_column or 1
        max_col = ws.max_column or 1
        report.append(f"  - **Satır Sayısı:** {max_row - min_row + 1}")
        report.append(f"  - **Sütun Sayısı:** {max_col - min_col + 1}")

        # Başlık satırını çıkar (ilk satır)
        headers = []
        try:
            for col in range(min_col, min(max_col + 1, 50)):  # max 50 sütun
                cell = ws.cell(row=min_row, column=col)
                if cell.value is not None:
                    headers.append(str(cell.value))
            if headers:
                report.append(f"  - **Başlık Sütunları:** {', '.join(f'`{h}`' for h in headers)}")
        except:
            pass

        # Birleştirilmiş hücreler
        if ws.merged_cells.ranges:
            report.append(f"  - **Birleştirilmiş Hücreler:** {len(ws.merged_cells.ranges)} adet")
            merged_list = [str(m) for m in list(ws.merged_cells.ranges)[:20]]
            report.append(f"    - Aralıklar: {', '.join(f'`{m}`' for m in merged_list)}")
            if len(ws.merged_cells.ranges) > 20:
                report.append(f"    - ... ve {len(ws.merged_cells.ranges) - 20} adet daha")

        # Formülleri bul
        formulas = []
        cell_dependencies = defaultdict(list)
        try:
            for row in ws.iter_rows(min_row=min_row, max_row=min(max_row, 5000),
                                     min_col=min_col, max_col=min(max_col, 100)):
                for cell in row:
                    if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                        cell_ref = f"{get_column_letter(cell.column)}{cell.row}"
                        formulas.append((cell_ref, cell.value))
                        # Bağımlılıkları çıkar
                        refs = re.findall(r"[A-Z]+[0-9]+(?::[A-Z]+[0-9]+)?", cell.value)
                        for ref in refs:
                            cell_dependencies[cell_ref].append(ref)
        except:
            pass

        if formulas:
            report.append(f"  - **Formül Sayısı:** {len(formulas)}")
            report.append("")
            report.append("  <details>")
            report.append(f"  <summary>📐 Formüller ({min(len(formulas), 30)} / {len(formulas)} gösteriliyor)</summary>")
            report.append("")
            report.append("  | Hücre | Formül |")
            report.append("  |-------|--------|")
            for cell_ref, formula in formulas[:30]:
                safe_formula = formula.replace('|', '\\|')
                report.append(f"  | `{cell_ref}` | `{safe_formula}` |")
            if len(formulas) > 30:
                report.append(f"  | ... | *{len(formulas) - 30} formül daha* |")
            report.append("")
            report.append("  </details>")
            report.append("")

        # Hücre bağımlılıkları
        if cell_dependencies:
            report.append("  <details>")
            report.append(f"  <summary>🔗 Hücre Bağımlılıkları ({min(len(cell_dependencies), 20)} / {len(cell_dependencies)} gösteriliyor)</summary>")
            report.append("")
            report.append("  | Hücre | Bağımlı Olduğu Hücreler |")
            report.append("  |-------|--------------------------|")
            for i, (cell_ref, deps) in enumerate(cell_dependencies.items()):
                if i >= 20:
                    break
                deps_str = ', '.join(f'`{d}`' for d in deps)
                report.append(f"  | `{cell_ref}` | {deps_str} |")
            if len(cell_dependencies) > 20:
                report.append(f"  | ... | *{len(cell_dependencies) - 20} bağımlılık daha* |")
            report.append("")
            report.append("  </details>")
            report.append("")

        # Tablolar (ListObject)
        if hasattr(ws, 'tables') and ws.tables:
            report.append(f"  - **Tablo Sayısı:** {len(ws.tables)}")
            for table_name, table in ws.tables.items():
                report.append(f"    - 📊 Tablo: `{table_name}` → Aralık: `{table.ref}`")
                if hasattr(table, 'tableColumns') and table.tableColumns:
                    cols = [tc.name for tc in table.tableColumns]
                    report.append(f"      - Sütunlar: {', '.join(f'`{c}`' for c in cols)}")

        # Veri doğrulama
        if hasattr(ws, 'data_validations') and ws.data_validations.dataValidation:
            report.append(f"  - **Veri Doğrulama Kuralları:** {len(ws.data_validations.dataValidation)} adet")
            for dv in ws.data_validations.dataValidation[:10]:
                report.append(f"    - Hücreler: `{dv.sqref}` | Tip: `{dv.type}` | Değerler: `{dv.formula1}`")

        # Koşullu biçimlendirme
        if ws.conditional_formatting:
            cf_count = len(list(ws.conditional_formatting))
            report.append(f"  - **Koşullu Biçimlendirme:** {cf_count} kural")

        # Grafikler
        if hasattr(ws, '_charts') and ws._charts:
            report.append(f"  - **Grafik Sayısı:** {len(ws._charts)}")
        
        # Yorumlar / Notlar
        comments = []
        try:
            for row in ws.iter_rows(min_row=min_row, max_row=min(max_row, 5000),
                                     min_col=min_col, max_col=min(max_col, 100)):
                for cell in row:
                    if cell.comment:
                        cell_ref = f"{get_column_letter(cell.column)}{cell.row}"
                        comments.append((cell_ref, cell.comment.text[:100]))
        except:
            pass
        
        if comments:
            report.append(f"  - **Yorum/Not Sayısı:** {len(comments)}")
            for cell_ref, text in comments[:10]:
                report.append(f"    - `{cell_ref}`: {text}")

        # Örnek veri (ilk 5 satır)
        report.append("")
        report.append("  **Örnek Veri (İlk 5 Satır):**")
        try:
            sample_rows = list(ws.iter_rows(
                min_row=min_row, max_row=min(min_row + 5, max_row + 1),
                min_col=min_col, max_col=min(max_col + 1, 15),
                values_only=True
            ))
            if sample_rows and any(any(c is not None for c in row) for row in sample_rows):
                # Markdown tablo oluştur
                col_count = len(sample_rows[0]) if sample_rows else 0
                if col_count > 0:
                    header_row = sample_rows[0] if sample_rows else []
                    h_cells = [str(c)[:30] if c is not None else '' for c in header_row]
                    report.append("  | " + " | ".join(h_cells) + " |")
                    report.append("  | " + " | ".join(['---'] * col_count) + " |")
                    for row in sample_rows[1:]:
                        cells = [str(c)[:30] if c is not None else '' for c in row]
                        report.append("  | " + " | ".join(cells) + " |")
            else:
                report.append("  > Sayfa boş veya veri bulunamadı.")
        except Exception as e:
            report.append(f"  > Örnek veri okunamadı: {e}")

        report.append("")

    wb.close()
    return "\n".join(report)


def _analyze_excel_xls(filepath):
    """Eski format .xls dosyasını analiz eder."""
    report = []
    if not HAS_XLRD:
        report.append("  > ⚠️ `.xls` formatı için `xlrd` kütüphanesi gerekli.")
        report.append("  > Yüklemek için: `pip install xlrd`")
        return "\n".join(report)
    
    try:
        wb = xlrd.open_workbook(filepath)
        report.append(f"  - **Sayfa Sayısı:** {wb.nsheets}")
        report.append(f"  - **Sayfalar:** {', '.join(wb.sheet_names())}")
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            report.append(f"  #### 📄 Sayfa: `{sheet_name}`")
            report.append(f"  - **Satır Sayısı:** {ws.nrows}")
            report.append(f"  - **Sütun Sayısı:** {ws.ncols}")
            # İlk satır (başlıklar)
            if ws.nrows > 0:
                headers = [str(ws.cell_value(0, c)) for c in range(min(ws.ncols, 50)) if ws.cell_value(0, c)]
                if headers:
                    report.append(f"  - **Başlık Sütunları:** {', '.join(f'`{h}`' for h in headers)}")
            report.append("")
    except Exception as e:
        report.append(f"  > ❌ .xls dosyası okunamadı: {e}")
    
    return "\n".join(report)


# ═══════════════════════════════════════════════════════════════════════════════
# WORD ANALİZİ
# ═══════════════════════════════════════════════════════════════════════════════
def analyze_word(filepath):
    """Word dosyasını analiz eder."""
    report = []
    
    if not HAS_DOCX:
        report.append("  > ⚠️ `python-docx` kütüphanesi yüklü değil. Word analizi yapılamadı.")
        report.append("  > Yüklemek için: `pip install python-docx`")
        return "\n".join(report)
    
    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        report.append(f"  > ❌ Dosya açılamadı: {e}")
        return "\n".join(report)
    
    # Genel bilgiler
    paragraphs = doc.paragraphs
    tables = doc.tables
    
    # Metin istatistikleri
    total_text = "\n".join([p.text for p in paragraphs if p.text.strip()])
    word_count = len(total_text.split()) if total_text else 0
    char_count = len(total_text)
    non_empty_paragraphs = [p for p in paragraphs if p.text.strip()]
    
    report.append(f"  - **Paragraf Sayısı:** {len(non_empty_paragraphs)}")
    report.append(f"  - **Kelime Sayısı:** {word_count}")
    report.append(f"  - **Karakter Sayısı:** {char_count}")
    report.append(f"  - **Tablo Sayısı:** {len(tables)}")
    
    # Bölüm bilgisi
    if doc.sections:
        report.append(f"  - **Bölüm (Section) Sayısı:** {len(doc.sections)}")
    
    # Başlıkları çıkar
    headings = []
    for p in paragraphs:
        if p.style and p.style.name and p.style.name.startswith('Heading'):
            level = p.style.name.replace('Heading', '').replace(' ', '')
            try:
                level_num = int(level)
            except:
                level_num = 1
            if p.text.strip():
                headings.append((level_num, p.text.strip()))
    
    if headings:
        report.append("")
        report.append("  #### 📑 Başlık Yapısı (İçindekiler)")
        for level, text in headings:
            indent = "  " * level
            report.append(f"  {indent}- {text}")
        report.append("")
    
    # Tabloları analiz et
    if tables:
        report.append("  #### 📊 Tablolar")
        for i, table in enumerate(tables[:10]):
            report.append(f"  **Tablo {i + 1}:** {len(table.rows)} satır × {len(table.columns)} sütun")
            # İlk satırı başlık olarak göster
            if table.rows:
                try:
                    header_cells = [cell.text.strip()[:30] for cell in table.rows[0].cells]
                    if any(header_cells):
                        report.append("  | " + " | ".join(header_cells) + " |")
                        report.append("  | " + " | ".join(['---'] * len(header_cells)) + " |")
                        # İlk 3 veri satırı
                        for row in table.rows[1:4]:
                            cells = [cell.text.strip()[:30] for cell in row.cells]
                            report.append("  | " + " | ".join(cells) + " |")
                        if len(table.rows) > 4:
                            report.append(f"  | *... {len(table.rows) - 4} satır daha* | |")
                except:
                    pass
            report.append("")
        if len(tables) > 10:
            report.append(f"  > ... ve {len(tables) - 10} tablo daha")
    
    # Metin önizleme (ilk 500 karakter)
    if total_text:
        preview = total_text[:500].replace('\n', ' ').strip()
        report.append("")
        report.append("  #### 📝 İçerik Önizleme")
        report.append(f"  > {preview}{'...' if len(total_text) > 500 else ''}")
    
    # Resimler
    image_count = 0
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
    except:
        pass
    if image_count:
        report.append(f"  - **Gömülü Resim Sayısı:** {image_count}")
    
    # Üstbilgi / Altbilgi
    header_texts = []
    footer_texts = []
    try:
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        header_texts.append(p.text.strip())
            if section.footer and section.footer.paragraphs:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        footer_texts.append(p.text.strip())
    except:
        pass
    
    if header_texts:
        report.append(f"  - **Üstbilgi:** {'; '.join(header_texts[:3])}")
    if footer_texts:
        report.append(f"  - **Altbilgi:** {'; '.join(footer_texts[:3])}")
    
    return "\n".join(report)


# ═══════════════════════════════════════════════════════════════════════════════
# PDF ANALİZİ
# ═══════════════════════════════════════════════════════════════════════════════
def analyze_pdf(filepath):
    """PDF dosyasını analiz eder."""
    report = []
    
    if not HAS_PDF:
        report.append("  > ⚠️ `pdfplumber` kütüphanesi yüklü değil. PDF analizi yapılamadı.")
        report.append("  > Yüklemek için: `pip install pdfplumber`")
        return "\n".join(report)
    
    try:
        with pdfplumber.open(filepath) as pdf:
            page_count = len(pdf.pages)
            report.append(f"  - **Sayfa Sayısı:** {page_count}")
            
            # Metadata
            if pdf.metadata:
                meta = pdf.metadata
                if meta.get('Title'):
                    report.append(f"  - **Başlık:** {meta['Title']}")
                if meta.get('Author'):
                    report.append(f"  - **Yazar:** {meta['Author']}")
                if meta.get('Subject'):
                    report.append(f"  - **Konu:** {meta['Subject']}")
                if meta.get('Creator'):
                    report.append(f"  - **Oluşturan:** {meta['Creator']}")
                if meta.get('CreationDate'):
                    report.append(f"  - **Oluşturma Tarihi:** {meta['CreationDate']}")
            
            # Her sayfayı analiz et
            all_text = []
            total_tables = 0
            total_images = 0
            
            report.append("")
            report.append("  #### 📄 Sayfa Detayları")
            
            for i, page in enumerate(pdf.pages[:50]):  # Max 50 sayfa detay
                page_text = page.extract_text() or ""
                word_count = len(page_text.split())
                tables = page.extract_tables() or []
                images = page.images or []
                
                total_tables += len(tables)
                total_images += len(images)
                all_text.append(page_text)
                
                if i < 10:  # İlk 10 sayfa detay göster
                    report.append(f"  - **Sayfa {i + 1}:** {word_count} kelime"
                                  f"{f', {len(tables)} tablo' if tables else ''}"
                                  f"{f', {len(images)} resim' if images else ''}")
            
            report.append("")
            report.append(f"  - **Toplam Tablo Sayısı:** {total_tables}")
            report.append(f"  - **Toplam Resim Sayısı:** {total_images}")
            
            # İlk sayfanın tablolarını göster
            for i, page in enumerate(pdf.pages[:5]):
                tables = page.extract_tables() or []
                for j, table in enumerate(tables[:3]):
                    if table and len(table) > 0:
                        report.append(f"")
                        report.append(f"  **Sayfa {i+1} - Tablo {j+1}:**")
                        # Başlık satırı
                        if table[0]:
                            headers = [str(c)[:25] if c else '' for c in table[0]]
                            report.append("  | " + " | ".join(headers) + " |")
                            report.append("  | " + " | ".join(['---'] * len(headers)) + " |")
                            for row in table[1:4]:
                                cells = [str(c)[:25] if c else '' for c in row]
                                # Sütun sayısını eşitle
                                while len(cells) < len(headers):
                                    cells.append('')
                                report.append("  | " + " | ".join(cells[:len(headers)]) + " |")
                            if len(table) > 4:
                                report.append(f"  | *... {len(table) - 4} satır daha* |")
            
            # İçerik önizleme
            full_text = "\n".join(all_text)
            total_word_count = len(full_text.split())
            report.append(f"")
            report.append(f"  - **Toplam Kelime Sayısı:** {total_word_count}")
            
            if full_text.strip():
                preview = full_text[:500].replace('\n', ' ').strip()
                report.append("")
                report.append("  #### 📝 İçerik Önizleme")
                report.append(f"  > {preview}{'...' if len(full_text) > 500 else ''}")
    
    except Exception as e:
        report.append(f"  > ❌ PDF okunamadı: {e}")
    
    return "\n".join(report)


# ═══════════════════════════════════════════════════════════════════════════════
# KLASÖR ANALİZİ
# ═══════════════════════════════════════════════════════════════════════════════
def build_tree(root_path, prefix="", is_last=True, max_depth=6, current_depth=0):
    """Klasör ağacı oluşturur (metin formatında)."""
    root = Path(root_path)
    tree_lines = []
    
    if current_depth == 0:
        tree_lines.append(f"📁 {root.name}/")
    
    if current_depth >= max_depth:
        tree_lines.append(f"{prefix}{'└── ' if is_last else '├── '}... (derinlik sınırına ulaşıldı)")
        return tree_lines
    
    try:
        entries = sorted(root.iterdir(), key=lambda e: (not e.is_dir(), e.name.lower()))
    except PermissionError:
        return tree_lines
    
    # Filtreleme
    entries = [e for e in entries if e.name not in IGNORED_DIRS 
               and not e.name.startswith('.')
               and e.name != RAPOR_KLASOR_ADI
               and e.name != RAPOR_DOSYA_ADI]
    
    for i, entry in enumerate(entries):
        is_entry_last = (i == len(entries) - 1)
        connector = "└── " if is_entry_last else "├── "
        extension = "    " if is_entry_last else "│   "
        
        if entry.is_dir():
            tree_lines.append(f"{prefix}{connector}📁 {entry.name}/")
            subtree = build_tree(entry, prefix + extension, is_entry_last, max_depth, current_depth + 1)
            tree_lines.extend(subtree)
        else:
            size = format_size(entry.stat().st_size) if entry.exists() else "?"
            icon = get_file_icon(entry.suffix.lower())
            tree_lines.append(f"{prefix}{connector}{icon} {entry.name} ({size})")
    
    return tree_lines


def get_file_icon(ext):
    """Dosya uzantısına göre ikon döndürür."""
    if ext in EXCEL_EXTENSIONS or ext in EXCEL_OLD_EXTENSIONS:
        return "📊"
    elif ext in WORD_EXTENSIONS:
        return "📝"
    elif ext in PDF_EXTENSIONS:
        return "📕"
    elif ext in IMAGE_EXTENSIONS:
        return "🖼️"
    elif ext in CODE_EXTENSIONS:
        return "💻"
    elif ext in ARCHIVE_EXTENSIONS:
        return "📦"
    elif ext in {'.mp4', '.avi', '.mov', '.mkv', '.webm'}:
        return "🎬"
    elif ext in {'.mp3', '.wav', '.flac', '.aac', '.ogg'}:
        return "🎵"
    else:
        return "📄"


def categorize_files(folder_path):
    """Klasördeki dosyaları kategorilere ayırır."""
    categories = {
        'excel': [],
        'word': [],
        'pdf': [],
        'code': [],
        'image': [],
        'archive': [],
        'other': [],
    }
    
    folder = Path(folder_path)
    try:
        for item in folder.iterdir():
            if item.is_file() and not item.name.startswith('.') and item.name != RAPOR_DOSYA_ADI:
                ext = item.suffix.lower()
                if ext in EXCEL_EXTENSIONS or ext in EXCEL_OLD_EXTENSIONS:
                    categories['excel'].append(item)
                elif ext in WORD_EXTENSIONS:
                    categories['word'].append(item)
                elif ext in PDF_EXTENSIONS:
                    categories['pdf'].append(item)
                elif ext in CODE_EXTENSIONS:
                    categories['code'].append(item)
                elif ext in IMAGE_EXTENSIONS:
                    categories['image'].append(item)
                elif ext in ARCHIVE_EXTENSIONS:
                    categories['archive'].append(item)
                else:
                    categories['other'].append(item)
    except PermissionError:
        pass
    
    return categories


# ═══════════════════════════════════════════════════════════════════════════════
# RAPOR OLUŞTURMA
# ═══════════════════════════════════════════════════════════════════════════════
def generate_folder_report(folder_path, root_path):
    """Bir klasör için detaylı MD rapor oluşturur."""
    folder = Path(folder_path)
    root = Path(root_path)
    relative = folder.relative_to(root)
    
    report = []
    report.append(f"# 📁 Klasör Raporu: `{folder.name}`")
    report.append(f"")
    report.append(f"**Tam Yol:** `{relative}`  ")
    report.append(f"**Rapor Tarihi:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append(f"")
    report.append("---")
    report.append("")
    
    # Dosya kategorileri
    categories = categorize_files(folder_path)
    
    # Alt klasörler
    subdirs = [d for d in folder.iterdir() if d.is_dir() 
               and d.name not in IGNORED_DIRS 
               and not d.name.startswith('.')
               and d.name != RAPOR_KLASOR_ADI]
    
    # Özet tablo
    total_files = sum(len(v) for v in categories.values())
    report.append("## 📋 Özet")
    report.append("")
    report.append(f"| Öğe | Sayı |")
    report.append(f"|-----|------|")
    report.append(f"| Alt Klasörler | {len(subdirs)} |")
    report.append(f"| Toplam Dosya | {total_files} |")
    if categories['excel']:
        report.append(f"| Excel Dosyaları | {len(categories['excel'])} |")
    if categories['word']:
        report.append(f"| Word Dosyaları | {len(categories['word'])} |")
    if categories['pdf']:
        report.append(f"| PDF Dosyaları | {len(categories['pdf'])} |")
    if categories['code']:
        report.append(f"| Kod Dosyaları | {len(categories['code'])} |")
    if categories['image']:
        report.append(f"| Resim Dosyaları | {len(categories['image'])} |")
    if categories['archive']:
        report.append(f"| Arşiv Dosyaları | {len(categories['archive'])} |")
    if categories['other']:
        report.append(f"| Diğer Dosyalar | {len(categories['other'])} |")
    report.append("")
    
    # Alt klasörler listesi
    if subdirs:
        report.append("## 📂 Alt Klasörler")
        report.append("")
        for d in sorted(subdirs, key=lambda x: x.name.lower()):
            sub_file_count = sum(1 for _ in d.rglob('*') if _.is_file())
            report.append(f"- 📁 **{d.name}/** — {sub_file_count} dosya")
        report.append("")
    
    # Dosya listesi
    if total_files > 0:
        report.append("## 📄 Dosya Listesi")
        report.append("")
        report.append("| Dosya | Tür | Boyut | Son Değişiklik |")
        report.append("|-------|-----|-------|----------------|")
        all_files = []
        for cat_files in categories.values():
            all_files.extend(cat_files)
        for f in sorted(all_files, key=lambda x: x.name.lower()):
            try:
                stat = f.stat()
                size = format_size(stat.st_size)
                modified = datetime.datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M')
                icon = get_file_icon(f.suffix.lower())
                report.append(f"| {icon} `{f.name}` | `{f.suffix}` | {size} | {modified} |")
            except:
                report.append(f"| `{f.name}` | `{f.suffix}` | ? | ? |")
        report.append("")
    
    # Excel analizi
    if categories['excel']:
        report.append("---")
        report.append("")
        report.append("## 📊 Excel Dosya Analizleri")
        report.append("")
        for excel_file in sorted(categories['excel'], key=lambda x: x.name.lower()):
            report.append(f"### 📊 `{excel_file.name}`")
            report.append(f"**Boyut:** {format_size(excel_file.stat().st_size)}")
            report.append("")
            try:
                analysis = analyze_excel(str(excel_file))
                report.append(analysis)
            except Exception as e:
                report.append(f"  > ❌ Excel analiz hatası: {e}")
            report.append("")
    
    # Word analizi
    if categories['word']:
        report.append("---")
        report.append("")
        report.append("## 📝 Word Dosya Analizleri")
        report.append("")
        for word_file in sorted(categories['word'], key=lambda x: x.name.lower()):
            report.append(f"### 📝 `{word_file.name}`")
            report.append(f"**Boyut:** {format_size(word_file.stat().st_size)}")
            report.append("")
            try:
                analysis = analyze_word(str(word_file))
                report.append(analysis)
            except Exception as e:
                report.append(f"  > ❌ Word analiz hatası: {e}")
            report.append("")
    
    # PDF analizi
    if categories['pdf']:
        report.append("---")
        report.append("")
        report.append("## 📕 PDF Dosya Analizleri")
        report.append("")
        for pdf_file in sorted(categories['pdf'], key=lambda x: x.name.lower()):
            report.append(f"### 📕 `{pdf_file.name}`")
            report.append(f"**Boyut:** {format_size(pdf_file.stat().st_size)}")
            report.append("")
            try:
                analysis = analyze_pdf(str(pdf_file))
                report.append(analysis)
            except Exception as e:
                report.append(f"  > ❌ PDF analiz hatası: {e}")
            report.append("")
    
    # Kod dosyaları
    if categories['code']:
        report.append("---")
        report.append("")
        report.append("## 💻 Kod Dosyaları")
        report.append("")
        for code_file in sorted(categories['code'], key=lambda x: x.name.lower()):
            report.append(f"- 💻 `{code_file.name}` ({format_size(code_file.stat().st_size)})")
            # Kısa açıklama için ilk birkaç satırı oku
            try:
                with open(code_file, 'r', encoding='utf-8', errors='ignore') as f:
                    first_lines = []
                    for line_num, line in enumerate(f):
                        if line_num >= 5:
                            break
                        stripped = line.strip()
                        if stripped and not stripped.startswith('#!'):
                            first_lines.append(stripped)
                    if first_lines:
                        # Docstring veya yorum varsa göster
                        for line in first_lines[:3]:
                            if line.startswith(('#', '//', '/*', '"""', "'''", '*')):
                                report.append(f"  > {line[:80]}")
                                break
            except:
                pass
        report.append("")
    
    # Resimler
    if categories['image']:
        report.append("---")
        report.append("")
        report.append("## 🖼️ Resim Dosyaları")
        report.append("")
        for img in sorted(categories['image'], key=lambda x: x.name.lower()):
            report.append(f"- 🖼️ `{img.name}` ({format_size(img.stat().st_size)})")
        report.append("")
    
    # Arşivler
    if categories['archive']:
        report.append("---")
        report.append("")
        report.append("## 📦 Arşiv Dosyaları")
        report.append("")
        for arc in sorted(categories['archive'], key=lambda x: x.name.lower()):
            report.append(f"- 📦 `{arc.name}` ({format_size(arc.stat().st_size)})")
        report.append("")
    
    # Diğer
    if categories['other']:
        report.append("---")
        report.append("")
        report.append("## 📄 Diğer Dosyalar")
        report.append("")
        for other in sorted(categories['other'], key=lambda x: x.name.lower()):
            report.append(f"- 📄 `{other.name}` ({format_size(other.stat().st_size)})")
        report.append("")
    
    return "\n".join(report)


def generate_master_report(root_path, folder_reports):
    """Ana rapor dosyasını oluşturur."""
    root = Path(root_path)
    
    report = []
    report.append(f"# 🏗️ Proje Analiz Raporu")
    report.append(f"")
    report.append(f"**Proje Kök Klasörü:** `{root.name}`  ")
    report.append(f"**Tam Yol:** `{root.resolve()}`  ")
    report.append(f"**Rapor Tarihi:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  ")
    report.append(f"**Rapor Oluşturan:** Proje Klasör Analiz Scripti v1.0")
    report.append(f"")
    report.append("---")
    report.append("")
    
    # Genel istatistikler
    total_files = 0
    total_dirs = 0
    total_size = 0
    file_type_stats = defaultdict(int)
    
    for dirpath, dirnames, filenames in os.walk(root_path):
        # Yoksayılan klasörleri filtrele
        dirnames[:] = [d for d in dirnames if d not in IGNORED_DIRS 
                       and not d.startswith('.') and d != RAPOR_KLASOR_ADI]
        total_dirs += len(dirnames)
        for fname in filenames:
            if not fname.startswith('.') and fname != RAPOR_DOSYA_ADI:
                total_files += 1
                fpath = os.path.join(dirpath, fname)
                try:
                    total_size += os.path.getsize(fpath)
                except:
                    pass
                ext = Path(fname).suffix.lower()
                if ext:
                    file_type_stats[ext] += 1
    
    report.append("## 📊 Genel İstatistikler")
    report.append("")
    report.append(f"| Metrik | Değer |")
    report.append(f"|--------|-------|")
    report.append(f"| Toplam Klasör | {total_dirs} |")
    report.append(f"| Toplam Dosya | {total_files} |")
    report.append(f"| Toplam Boyut | {format_size(total_size)} |")
    report.append(f"| Analiz Edilen Klasör | {len(folder_reports)} |")
    report.append("")
    
    # Dosya türü dağılımı
    if file_type_stats:
        report.append("## 📈 Dosya Türü Dağılımı")
        report.append("")
        report.append("| Uzantı | Sayı | Kategori |")
        report.append("|--------|------|----------|")
        for ext, count in sorted(file_type_stats.items(), key=lambda x: -x[1]):
            cat = "📊 Excel" if ext in EXCEL_EXTENSIONS | EXCEL_OLD_EXTENSIONS else \
                  "📝 Word" if ext in WORD_EXTENSIONS else \
                  "📕 PDF" if ext in PDF_EXTENSIONS else \
                  "💻 Kod" if ext in CODE_EXTENSIONS else \
                  "🖼️ Resim" if ext in IMAGE_EXTENSIONS else \
                  "📦 Arşiv" if ext in ARCHIVE_EXTENSIONS else "📄 Diğer"
            report.append(f"| `{ext}` | {count} | {cat} |")
        report.append("")
    
    # Klasör ağacı
    report.append("## 🌳 Klasör Yapısı")
    report.append("")
    report.append("```")
    tree = build_tree(root_path)
    report.append("\n".join(tree))
    report.append("```")
    report.append("")
    
    # Klasör raporlarına bağlantılar
    report.append("## 📂 Klasör Raporları")
    report.append("")
    report.append("Her klasörün detaylı raporu hem kendi içinde hem de bu rapor klasöründe bulunabilir.")
    report.append("")
    report.append("| Klasör | Dosya Sayısı | Rapor |")
    report.append("|--------|-------------|-------|")
    
    for folder_path, report_filename in sorted(folder_reports, key=lambda x: x[0]):
        folder = Path(folder_path)
        try:
            relative = folder.relative_to(root)
        except:
            relative = folder
        
        file_count = sum(1 for f in folder.iterdir() if f.is_file() 
                        and not f.name.startswith('.') and f.name != RAPOR_DOSYA_ADI) \
                    if folder.exists() else 0
        
        report.append(f"| 📁 `{relative}` | {file_count} | [{report_filename}]({report_filename}) |")
    
    report.append("")
    
    # Her klasörün raporunu da ana rapora ekle
    report.append("---")
    report.append("")
    report.append("# 📑 Detaylı Klasör Raporları")
    report.append("")
    
    return "\n".join(report)


# ═══════════════════════════════════════════════════════════════════════════════
# ANA İŞLEM
# ═══════════════════════════════════════════════════════════════════════════════
def run_analysis(target_path):
    """Ana analiz işlemini çalıştırır."""
    root = Path(target_path).resolve()
    
    if not root.exists():
        print(f"❌ Hata: '{root}' yolu bulunamadı!")
        sys.exit(1)
    
    if not root.is_dir():
        print(f"❌ Hata: '{root}' bir klasör değil!")
        sys.exit(1)
    
    print(f"")
    print(f"╔══════════════════════════════════════════════════════╗")
    print(f"║       🔍 Proje Klasör Analiz Scripti v1.0          ║")
    print(f"╚══════════════════════════════════════════════════════╝")
    print(f"")
    print(f"📁 Hedef: {root}")
    print(f"📅 Tarih: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"")
    
    # Kütüphane durumu
    print("📦 Kütüphane Durumu:")
    print(f"   {'✅' if HAS_OPENPYXL else '❌'} openpyxl (Excel .xlsx)")
    print(f"   {'✅' if HAS_XLRD else '❌'} xlrd (Excel .xls)")
    print(f"   {'✅' if HAS_DOCX else '❌'} python-docx (Word)")
    print(f"   {'✅' if HAS_PDF else '❌'} pdfplumber (PDF)")
    
    missing = []
    if not HAS_OPENPYXL: missing.append('openpyxl')
    if not HAS_DOCX: missing.append('python-docx')
    if not HAS_PDF: missing.append('pdfplumber')
    if not HAS_XLRD: missing.append('xlrd')
    
    if missing:
        print(f"\n⚠️  Eksik kütüphaneler: {', '.join(missing)}")
        print(f"   Yüklemek için: pip install {' '.join(missing)}")
    
    print(f"\n{'─' * 55}")
    print(f"🔄 Analiz başlıyor...\n")
    
    # Rapor klasörü oluştur
    report_dir = root / RAPOR_KLASOR_ADI
    report_dir.mkdir(exist_ok=True)
    
    # Tüm klasörleri tara
    folders_to_analyze = []
    for dirpath, dirnames, filenames in os.walk(str(root)):
        # Yoksayılan klasörleri filtrele
        dirnames[:] = [d for d in dirnames if d not in IGNORED_DIRS 
                       and not d.startswith('.') and d != RAPOR_KLASOR_ADI]
        
        # Dosya var mı kontrol et (rapor dosyaları hariç)
        real_files = [f for f in filenames if not f.startswith('.') and f != RAPOR_DOSYA_ADI]
        if real_files or dirpath == str(root):
            folders_to_analyze.append(dirpath)
    
    print(f"📂 {len(folders_to_analyze)} klasör analiz edilecek.\n")
    
    folder_reports = []
    
    for i, folder_path in enumerate(folders_to_analyze):
        folder = Path(folder_path)
        try:
            relative = folder.relative_to(root)
        except:
            relative = folder
        
        print(f"  [{i+1}/{len(folders_to_analyze)}] 📁 {relative or '.'} ...", end=" ", flush=True)
        
        # Klasör raporu oluştur
        report_content = generate_folder_report(folder_path, str(root))
        
        # Raporu klasörün içine kaydet
        local_report_path = folder / RAPOR_DOSYA_ADI
        try:
            with open(local_report_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
        except Exception as e:
            print(f"⚠️ Yerel rapor yazılamadı: {e}")
        
        # Raporu ana rapor klasörüne de kaydet
        if str(relative) == '.':
            report_filename = "KOK_KLASOR_RAPORU.md"
        else:
            report_filename = str(relative).replace(os.sep, '_').replace('/', '_') + "_RAPORU.md"
        
        master_report_path = report_dir / report_filename
        try:
            with open(master_report_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
        except Exception as e:
            print(f"⚠️ Ana rapor yazılamadı: {e}")
        
        folder_reports.append((folder_path, report_filename))
        
        # Analiz edilen dosya sayısı
        categories = categorize_files(folder_path)
        analyzed = len(categories['excel']) + len(categories['word']) + len(categories['pdf'])
        total = sum(len(v) for v in categories.values())
        print(f"✅ ({total} dosya, {analyzed} analiz edildi)")
    
    # Ana rapor oluştur
    print(f"\n{'─' * 55}")
    print(f"📝 Ana rapor oluşturuluyor...")
    
    master_content = generate_master_report(str(root), folder_reports)
    
    # Her klasör raporunun içeriğini ana rapora ekle
    for folder_path, report_filename in sorted(folder_reports, key=lambda x: x[0]):
        folder = Path(folder_path)
        try:
            relative = folder.relative_to(root)
        except:
            relative = folder
        
        master_report_file = report_dir / report_filename
        if master_report_file.exists():
            with open(master_report_file, 'r', encoding='utf-8') as f:
                folder_content = f.read()
            master_content += f"\n---\n\n{folder_content}\n"
    
    # Ana raporu kaydet
    master_report_path = report_dir / "ANA_RAPOR.md"
    with open(master_report_path, 'w', encoding='utf-8') as f:
        f.write(master_content)
    
    # Ayrıca kök dizine bir özet rapor koy
    root_summary_path = root / "PROJE_ANALIZ_RAPORU.md"
    with open(root_summary_path, 'w', encoding='utf-8') as f:
        f.write(master_content)
    
    print(f"")
    print(f"╔══════════════════════════════════════════════════════╗")
    print(f"║              ✅ ANALİZ TAMAMLANDI!                  ║")
    print(f"╚══════════════════════════════════════════════════════╝")
    print(f"")
    print(f"📊 Sonuçlar:")
    print(f"   📁 Analiz edilen klasör: {len(folders_to_analyze)}")
    print(f"   📄 Oluşturulan rapor:    {len(folder_reports) + 1}")
    print(f"")
    print(f"📂 Rapor Konumları:")
    print(f"   🏠 Ana Rapor:        {root_summary_path}")
    print(f"   📁 Rapor Klasörü:    {report_dir}/")
    print(f"   📄 Klasör Raporları: Her klasörde '{RAPOR_DOSYA_ADI}'")
    print(f"")


# ═══════════════════════════════════════════════════════════════════════════════
# GİRİŞ NOKTASI
# ═══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Proje klasörlerini analiz edip MD raporları oluşturur.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Örnekler:
  python dosya_analiz.py /path/to/project
  python dosya_analiz.py .
  python dosya_analiz.py ~/Documents/MyProject

Bu script şunları yapar:
  • Klasör yapısını ağaç olarak çıkarır
  • Excel dosyalarını analiz eder (sayfalar, formüller, bağımlılıklar)
  • Word dosyalarını okur (başlıklar, tablolar, içerik)
  • PDF dosyalarını okur (sayfalar, tablolar, metin)
  • Her klasöre _KLASOR_RAPORU.md dosyası oluşturur
  • _ANALIZ_RAPORLARI/ klasöründe tüm raporları toplar
        """
    )
    parser.add_argument(
        "hedef",
        nargs="?",
        default=".",
        help="Analiz edilecek klasör yolu (varsayılan: mevcut dizin)"
    )
    
    args = parser.parse_args()
    target = os.path.abspath(args.hedef)
    
    run_analysis(target)
